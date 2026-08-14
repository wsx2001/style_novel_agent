# backend/app/services/export/exporter.py
"""项目导出（docs/TECH.md §5.6）。

将项目所有章节导出为 txt / markdown / json / docx，写入 DATA_DIR/exports/。
- txt：纯文本（章节标题 + 正文）
- markdown：合并为一个 Markdown 文件（项目标题 + 章节）
- json：结构化数据（项目元数据 + 章节数组）
- docx：使用 python-docx 生成（标题 + 段落）
"""
from __future__ import annotations

import json
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Literal

from ...config import settings
from ...models import Chapter, Project

ExportFormat = Literal["txt", "markdown", "json", "docx"]

_EXTENSIONS: dict[ExportFormat, str] = {
    "txt": "txt",
    "markdown": "md",
    "json": "json",
    "docx": "docx",
}

# 文件名中 Windows 不允许的字符 → 下划线
_INVALID_CHARS = re.compile(r'[\\/:*?"<>|\s]+')


def _safe_filename(title: str) -> str:
    """清理文件名中的非法字符，空串回退为 project。"""
    cleaned = _INVALID_CHARS.sub("_", title).strip("_")
    return cleaned or "project"


def _render_txt(project: Project, chapters: list[Chapter]) -> str:
    parts = [f"《{project.title}》"]
    if project.genre:
        parts.append(f"类型：{project.genre}")
    if project.description:
        parts.append(project.description)
    parts.append("")
    for chapter in chapters:
        parts.append(f"## {chapter.title}")
        parts.append(chapter.content or "")
        parts.append("")
    return "\n".join(parts)


def _render_markdown(project: Project, chapters: list[Chapter]) -> str:
    lines = [f"# {project.title}", ""]
    if project.description:
        lines += [project.description, ""]
    for chapter in chapters:
        lines.append(f"## {chapter.title}")
        lines.append("")
        lines.append(chapter.content or "")
        lines.append("")
    return "\n".join(lines)


def _render_json(project: Project, chapters: list[Chapter]) -> str:
    payload = {
        "project": {
            "id": project.id,
            "title": project.title,
            "description": project.description,
            "genre": project.genre,
            "created_at": project.created_at.isoformat() if project.created_at else None,
            "updated_at": project.updated_at.isoformat() if project.updated_at else None,
        },
        "chapters": [
            {
                "id": chapter.id,
                "title": chapter.title,
                "order": chapter.order,
                "status": chapter.status,
                "content": chapter.content,
                "word_count": chapter.word_count,
                "created_at": chapter.created_at.isoformat() if chapter.created_at else None,
                "updated_at": chapter.updated_at.isoformat() if chapter.updated_at else None,
            }
            for chapter in chapters
        ],
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


def _render_docx(project: Project, chapters: list[Chapter]) -> bytes:
    """用 python-docx 生成 docx（延迟导入，避免影响其他格式）。"""
    from docx import Document

    doc = Document()
    doc.add_heading(project.title or "未命名项目", level=0)
    if project.genre:
        doc.add_paragraph(f"类型：{project.genre}")
    if project.description:
        doc.add_paragraph(project.description)
    for chapter in chapters:
        doc.add_heading(chapter.title, level=1)
        for line in (chapter.content or "").split("\n"):
            if line.strip():
                doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_project(
    project: Project, chapters: list[Chapter], format_: ExportFormat
) -> Path:
    """生成导出文件并写入 DATA_DIR/exports/，返回文件路径。"""
    exports_dir = settings.data_dir / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)
    date_str = datetime.utcnow().strftime("%Y%m%d")
    filename = f"{_safe_filename(project.title)}_{date_str}.{_EXTENSIONS[format_]}"
    path = exports_dir / filename

    if format_ == "txt":
        path.write_text(_render_txt(project, chapters), encoding="utf-8")
    elif format_ == "markdown":
        path.write_text(_render_markdown(project, chapters), encoding="utf-8")
    elif format_ == "json":
        path.write_text(_render_json(project, chapters), encoding="utf-8")
    else:  # docx
        path.write_bytes(_render_docx(project, chapters))
    return path
